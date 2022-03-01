from sanic import Sanic
from sanic import response
from sanic.log import logger
import sanic
import datetime
import random
import string
from Model2Server import Model2Server
import os
from sanic.response import redirect, html
from docx import Document
import json

random.seed(1003)

"""
这个函数的作用是产生并输出一个长度为length的随机字符串，
字符是从a-z,A-Z里随机选的，但我不太知道这是要干啥
"""
def generate_conv_id(length=10):
    letters = string.ascii_letters
    result_str = ''.join(random.choice(letters) for i in range(length))
    print("Random string of length", length, "is:", result_str)
    return result_str

global new_case_conv_id
new_case_conv_id = generate_conv_id()

conv_dict = {}
m2s = Model2Server()

with open('./data/item_info.json', 'r', encoding='utf-8') as f:
    item_info = json.load(f)

with open('./data/zh_index2attribute.json', 'r', encoding='utf-8') as f:
    zh_index2attr = json.load(f)

with open('./data/attribute_tree_dict.json', 'r', encoding='utf-8') as f:
    attr_tree = json.load(f)

def new_conv_id_init(conv_id):
    conv_dict[conv_id] = {}
    conv_dict[conv_id]['input_case'] = None
    conv_dict[conv_id]['pos_attribute_list'] = []
    conv_dict[conv_id]['neg_attribute_list'] = []
    conv_dict[conv_id]['parent_attribute_list'] = []
    conv_dict[conv_id]['rejected_item_list'] = []
    conv_dict[conv_id]['last_asked_parent_att'] = None
    conv_dict[conv_id]['last_asked_att_id_list'] = None
    conv_dict[conv_id]['last_rec_item_list'] = []
    conv_dict[conv_id]['last_system_action'] = None

#input_case, rejected_item_list
def set_input_case(conv_id, input_case):
    conv_dict[conv_id]['input_case'] = input_case
    conv_dict[conv_id]['rejected_item_list'] = [input_case]


#类似上一个函数
def set_input_new_case(conv_id, input_content):
    conv_dict[conv_id]['input_case'] = 1649  # or -1 ?
    duplicate_case_id = m2s.add_new_case(input_content)
    print("********duplicate_case_id********: ",duplicate_case_id)
    # print("duplicate_case_id: ", duplicate_case_id))
    if duplicate_case_id is not None:
        conv_dict[conv_id]['rejected_item_list'] = [duplicate_case_id]

#pos_attribute_list,neg_attribute_list,parent_attribute_list
def add_att(conv_id, pos_att, neg_att, parent_att):
    conv_dict[conv_id]['pos_attribute_list'].append(pos_att)
    conv_dict[conv_id]['neg_attribute_list'].append(neg_att)
    conv_dict[conv_id]['parent_attribute_list'].append(parent_att)

#rejected_item_list
def add_rejected_item(conv_id, rejected_item_list):
    for _ in rejected_item_list:
        if _ != 'null':
            conv_dict[conv_id]['rejected_item_list'].append(_)

#
def user_return_text(conv_id, user_text):
    print("user_text: ", user_text)
    if conv_dict[conv_id]['input_case'] is None:
        input_case = int(user_text)
        set_input_case(conv_id, input_case)
        # conv_dict[conv_id]['input_case'] = input_case
    else:

        user_text = user_text.strip().split(',')
        print("-------------user_text-----------: ", user_text)

        if user_text[0] == 'no' and conv_dict[conv_id]['last_system_action'] == 'rec':
            return

        pos_att = []
        if user_text[0] != 'no':
            pos_att = [int(_) for _ in user_text]

        asked_parent_att = conv_dict[conv_id]['last_asked_parent_att']
        asked_att_list = conv_dict[conv_id]['last_asked_att_id_list']
        neg_att = []
        for _ in asked_att_list:
            if _ not in pos_att:
                neg_att.append(_)

        add_att(conv_id, pos_att, neg_att, asked_parent_att)

#删除标签的操作，改变了pos_attribute_list,neg_attribute_list,parent_attribute_list
def delete_span(conv_id, user_text):
        # if user_text[0] != 'no':
        #     neg_att = [int(_) for _ in user_text]
        neg_att = [user_text]
        # print("pos:", conv_dict[conv_id]['pos_attribute_list'])
        # print("delete:", neg_att)

        for key, item in attr_tree.items():
            if neg_att[0] in item:
                parent = key

        asked_parent_att = parent
        asked_att_list = attr_tree[parent]
        ori_neg_att = []
        for _ in asked_att_list:  # 获取要删除的原neg_list
            if _ not in neg_att:
                ori_neg_att.append(_)

        print("pos list:", conv_dict[conv_id]['pos_attribute_list'])
        print("remove:", neg_att)
        print("neg list:", conv_dict[conv_id]['neg_attribute_list'])
        print("remove:", ori_neg_att)
        print("par list:", conv_dict[conv_id]['parent_attribute_list'])
        print("remove:", int(parent))

        conv_dict[conv_id]['pos_attribute_list'].remove(neg_att)
        conv_dict[conv_id]['neg_attribute_list'].remove(ori_neg_att)
        print("pos list2:", conv_dict[conv_id]['pos_attribute_list'])
        print("neg list2:", conv_dict[conv_id]['neg_attribute_list'])
        conv_dict[conv_id]['parent_attribute_list'].remove(int(parent))
        # conv_dict[conv_id]['neg_attribute_list'].append(asked_att_list)
        print("par list2:", conv_dict[conv_id]['parent_attribute_list'])

def get_recommeded_item(conv_id):
    input_case = conv_dict[conv_id]['input_case']
    print("get_recommeded_item!!!")
    return_dict = m2s.get_action_from_conv_his(conv_dict[conv_id], return_action=False, return_rec_items=True)
    id_list = return_dict['rec_item_id_list']
    print("id_list: ", id_list)
    text_list = return_dict['rec_item_text_list']
    print("text_list: ",text_list)
    text_mark_list = return_dict['rec_item_text_mark_list']
    print("text_mark_list: ", text_mark_list)
    # input_select_list = return_dict['input_item_text_select_list']
    # item_select_list = return_dict['rec_item_text_select_list']
    # print("text_mark_list: ", text_mark_list)
    conv_dict[conv_id]['last_rec_item_list'] = id_list

    assert len(id_list) == 5
    json_dict = {}
    if input_case == 1649:
        with open('new_case_text.json', 'r') as f:
            new_case_text = json.load(f)
        json_dict['rightwords1'] = f'上传案例: {new_case_text[str(input_case)]}'
    else:
        json_dict['rightwords1'] = f'上传案例: {m2s.dm.index2text[input_case]}'
    # print("id_list: ", id_list)
    # print("text_list: ", text_list)
    for idx, (text_id, text_content) in enumerate(zip(id_list, text_list)):
        # print("text_id, text_content: ", text_id, text_content)
        print("idx:",idx)
        print("text_id, text_content: ", text_id, text_content)
        if text_id != 'null':
            # f'rec case id {text_id}: {text_content}'
            # json_dict['rightwords' + str(idx + 2)] = f'相似案例{idx+1}: {text_content}'
            json_dict['rightwords' + str(idx + 2)] = f'{text_content}'
        else:
            json_dict['rightwords' + str(idx + 2)] = '没有更多啦~'

        ## 获取推荐item的feature
        # json_dict['keywords' + str(idx + 2)] = zh_index2attr[item_info[str(text_id)]]
        if text_id != 'null':
            key_list = []
            for attr in item_info[str(text_id)]:
                key_list.append(zh_index2attr[str(attr)])
            json_dict['keywords' + str(idx + 2)] = key_list
            json_dict['keylen' + str(idx + 2)] = len(key_list)
        else:
            json_dict['keywords' + str(idx + 2)] = []
            json_dict['keylen' + str(idx + 2)] = 0

    json_dict['mark_list'] = text_mark_list
    json_dict["input_select_list"] = return_dict['input_item_text_select_list']
    json_dict["item_select_list"] = return_dict['rec_item_text_select_list']

    print("key of json_dict: ", json_dict.keys())
    print("rightwords1: ",json_dict['rightwords1'])
    print("rightwords2: ", json_dict['rightwords2'])
    print("keywords2: ", json_dict['keywords2'])
    print("keylen2: ", json_dict['keylen2'])
    print("rightwords3: ", json_dict['rightwords3'])
    print("keywords3: ", json_dict['keywords3'])
    print("keylen3: ", json_dict['keylen3'])
    print("mark_list: ", json_dict['mark_list'])
    print("input_select_list: ", json_dict['input_select_list'])
    print("item_select_list: ", json_dict['item_select_list'])
    print("json_dict: ", json_dict)

    return json_dict


def user_return_continue(conv_id):
    add_rejected_item(conv_id, conv_dict[conv_id]['last_rec_item_list'])
    rec_item_dict = get_recommeded_item(conv_id)
    print("user_return_continue!!!")
    return_dict = m2s.get_action_from_conv_his(conv_dict[conv_id], return_action=True, return_rec_items=False)
    json_dict = {}
    print("return_dict['action_type']: ", return_dict['action_type'])
    if return_dict['action_type'] == 'rec':
        json_dict['agent_words'] = '已为您推荐一批案件  <br />'
        conv_dict[conv_id]['last_system_action'] = 'rec'
        # rec_item_dict = get_recommeded_item(conv_id)
    else:
        # json_dict['agent_words'] = 'choose attribute you need:  <br />  no for nothing  <br /> '
        json_dict['agent_words'] = '您想查询的案件有什么特征呢？  <br />'
        parent_att = return_dict['asked_parent_att']
        asked_att_id_list = return_dict['asked_att_id_list']
        asked_att_name_list = return_dict['asked_att_name_list']

        conv_dict[conv_id]['last_asked_parent_att'] = parent_att
        conv_dict[conv_id]['last_asked_att_id_list'] = asked_att_id_list
        conv_dict[conv_id]['last_system_action'] = 'ask'
        attribute_list_str = ' <br /> '.join(
            [f'{idx}: {att_name}' for idx, att_name in zip(asked_att_id_list, asked_att_name_list)])
        # json_dict['agent_words'] = json_dict['agent_words'] + attribute_list_str

        json_dict['attribute_list'] = asked_att_name_list
        json_dict['attribute_len'] = len(asked_att_name_list)
        json_dict['asked_att_id_list'] = asked_att_id_list

    # print("rec_item_dict: ", rec_item_dict.keys())
    json_dict.update(rec_item_dict)

    print("key of json_dict: ", json_dict.keys())
    print("agent_words：", json_dict['agent_words'])
    print("attribute_list：", json_dict['attribute_list'])
    print("attribute_len：", json_dict['attribute_len'])
    print("asked_att_id_list：", json_dict['asked_att_id_list'])
    print("json_dict：", json_dict)

    return json_dict


# 指定日志的文件名字，记录的日志记录在那个文件中
f_access = open("access.log", mode='a+', encoding='UTF-8')
f_error = open("error.log", mode='a+', encoding='UTF-8')
LOGGING_CONFIG = sanic.log.LOGGING_CONFIG_DEFAULTS
LOGGING_CONFIG['handlers']['error_console']['stream'] = f_error
LOGGING_CONFIG['handlers']['access_console']['stream'] = f_access

app = Sanic(name='ceshi', log_config=LOGGING_CONFIG)


# app.static('/js', './js')
# app.static('/css', './css')
app.static('/images', './images')
@app.route("/", methods=["POST"])
async def post_data(request):
    data_dict = request.json
    user_action_type = data_dict['user_action_type']

    if user_action_type == 'text':
        # print("user_action_type: ", user_action_type)
        conv_id = data_dict['conv_id']
        user_text = data_dict['user_text']
        user_return_text(conv_id, user_text)
        # return_dict = get_recommeded_item(conv_id)
        return_dict = {}
        # print(return_dict)

    elif user_action_type == 'reset':
        conv_id = generate_conv_id()
        global new_case_conv_id
        new_case_conv_id = conv_id
        new_conv_id_init(conv_id)
        return_dict = {}
        return_dict['conv_id'] = conv_id
        # print(return_dict)

    elif user_action_type == 'continue':
        if 'conv_id' in data_dict.keys():
            conv_id = data_dict['conv_id']
        else:
            conv_id = new_case_conv_id
        # print("conv_id2:", conv_id)
        return_dict = user_return_continue(conv_id)
        # print(return_dict)

    elif user_action_type == 'delete':
        # print("user_action_type: ", user_action_type)
        conv_id = data_dict['conv_id']
        user_text = data_dict['user_text']
        delete_span(conv_id, user_text)
        return_dict = get_recommeded_item(conv_id)
        # print(return_dict)

    return response.json(return_dict)

@app.route("/", methods=["GET"])
async def get_data(request):
    # print("get_data func")
    conv_id = new_case_conv_id
    new_conv_id_init(conv_id)
    html_content = open('ceshi.html', mode='r', encoding='utf8').read()
    html_content = html_content.replace('###########', conv_id)
    return response.html(html_content)

@app.route('/upload', methods=['POST', 'GET'])
async def upload(request):
    if request.method == 'POST':
        f = request.files.get("file")
        basepath = os.path.dirname(__file__)
        upload_path = os.path.join(basepath, 'static', f.name)
        # print(f.name)
        # print(f.type)
        # print(f.body)
        with open(upload_path, 'wb') as fw:
            fw.write(f.body)
        conv_id = new_case_conv_id
        new_conv_id_init(conv_id)
        user_text = 1649

        # print("upload_path:", upload_path)

        # document = Document(upload_path)
        # for paragraph in document.paragraphs:
        #     print("正文：\n", paragraph.text)
        #     set_input_new_case(conv_id, str(paragraph.text))
        #     text_dict = {}
        #     text_dict[user_text] = paragraph.text
        #     with open('new_case_text.json', 'w') as f:
        #         f.write(json.dumps(text_dict))
        # return_dict = get_recommeded_item(conv_id)
        # print(return_dict)

        document = Document(upload_path)
        document_content_list = []
        for paragraph in document.paragraphs:
            print("正文：\n", paragraph.text)
            document_content_list.append(str(paragraph.text))
        document_content = f'\n\n'.join([_ for _ in document_content_list if len(_)>0])

        set_input_new_case(conv_id, document_content)
        text_dict = {}
        text_dict[user_text] = document_content
        with open('new_case_text.json', 'w') as f:
            f.write(json.dumps(text_dict))
        # return_dict = get_recommeded_item(conv_id)
        return_dict = {}

    return response.json(return_dict)



# @app.route("/",methods=["POST"])
# async def post_data(request):
#     data_dict = request.json
#     print(data_dict)
#     return response.json(data_dict)
app.run(host='127.0.0.1', port=9999)
#app.run(host='183.174.228.109', port=1241)
f_access.close()
f_error.close()

print("1")